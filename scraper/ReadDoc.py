import docx


HEADERS = ['Education',
           'Relevant Exp',
           'Leadership and Act',
           'Leadership Exp',
           'Volunteer Exp',
           'Leadership & Act',
           'Professional Exp',
           'Skills and Int',
           'Skills & Int']

'''
def organize(l):
    exp_list = []
    in_count = 1
    out_count = 0
    for index in l:
        if in_count == 1:
            exp_list.append({})
            exp_list[out_count]['Organization'] = index
            in_count += 1
        elif in_count == 2:
            exp_list[out_count]['Location'] = index
            in_count += 1
        elif in_count == 3:
            exp_list[out_count]['Position'] = index
            in_count += 1
        elif in_count == 4:
            exp_list[out_count]['Dates'] = index
            in_count += 1
        elif in_count == 5:
            exp_list[out_count]['Bullet1'] = index
            in_count += 1
        elif in_count == 6:
            exp_list[out_count]['Bullet2'] = index
            out_count += 1
            in_count = 1
    return exp_list
'''


# Can only work if the there are two bullets per experience
def organize(l):
    exp_list = []
    in_count = 1
    out_count = 0
    for index in l:
        if in_count == 1:
            exp_list.append([])
            exp_list[out_count].append(index)
            in_count += 1
        elif in_count == 2:
            exp_list[out_count].append(index)
            in_count += 1
        elif in_count == 3:
            exp_list[out_count].append(index)
            in_count += 1
        elif in_count == 4:
            exp_list[out_count].append(index)
            in_count += 1
        elif in_count == 5:
            exp_list[out_count].append(index)
            in_count += 1
        elif in_count == 6:
            exp_list[out_count].append(index)
            out_count += 1
            in_count = 1
    return exp_list


def get_info(filename, header_string):
    doc = docx.Document(filename)
    exp_text = []
    start = False
    for para in doc.paragraphs:
        if not start:
            if para.text.find(header_string) != -1:
                start = True
        else:
            if not header(para):
                para_list = para.text.split('\t')
                for lines in para_list:
                    if lines != '':
                        exp_text.append(lines)
            else:
                break
    return exp_text


def get_experience(filename, header_string):
    return organize(get_info(filename, header_string))


def get_text(filename):
    doc = docx.Document(filename)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return full_text


def get_education(filename):
    return get_info(filename, 'Education')


def get_relevant_courses(filename):
    return get_info(filename, 'Relevant Exp')


def get_professional_experience(filename):
    return get_experience(filename, 'Professional Exp')


def get_leadership_experience(filename):
    return get_experience(filename, 'Leadership and Act')


def get_skills_interests(filename):
    return get_info(filename, 'Skills and Int')


def header(para):
    for h in HEADERS:
        if para.text.find(h) != -1:
            return True

    return False